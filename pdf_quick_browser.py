"""
PDF Quick Browser v3
Navigateur de fichiers avec preview dans une fenêtre séparée
- Fenêtre 1 : Explorateur de fichiers complet (arborescence + liste)
- Fenêtre 2 : Viewer multi-format (PDF, images, DOCX, RTF)
"""

import sys
import os
import re
import string
from datetime import datetime

from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QSplitter, QTreeView, QTreeWidget, QTreeWidgetItem, QHeaderView,
    QPushButton, QLabel, QFileDialog, QStatusBar, QLineEdit,
    QMenu, QAbstractItemView, QFrame, QToolButton, QSizePolicy,
    QFileSystemModel, QScrollArea, QTextBrowser, QStackedWidget
)
from PySide6.QtCore import (
    Qt, QTimer, QDir, QModelIndex, QSize, QFileInfo, Signal, QObject, QPointF
)
from PySide6.QtGui import (
    QKeySequence, QFont, QColor, QPalette, QShortcut,
    QIcon, QAction, QPixmap, QImage
)
from PySide6.QtPdf import QPdfDocument
from PySide6.QtPdfWidgets import QPdfView

# Optional imports for document formats
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from striprtf.striprtf import rtf_to_text
    HAS_RTF = True
except ImportError:
    HAS_RTF = False


# ─── Constants ───

PDF_EXTENSIONS = {'.pdf'}
IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.bmp', '.gif', '.webp', '.svg', '.ico', '.tiff', '.tif'}
DOCX_EXTENSIONS = {'.docx'}
RTF_EXTENSIONS = {'.rtf'}
PREVIEWABLE_EXTENSIONS = PDF_EXTENSIONS | IMAGE_EXTENSIONS | DOCX_EXTENSIONS | RTF_EXTENSIONS


# ─── Helpers ───

def format_size(size_bytes):
    if size_bytes < 1024:
        return f"{size_bytes} o"
    elif size_bytes < 1024 * 1024:
        return f"{size_bytes / 1024:.1f} Ko"
    else:
        return f"{size_bytes / (1024 * 1024):.1f} Mo"


def format_date(timestamp):
    dt = datetime.fromtimestamp(timestamp)
    months = [
        "", "janv.", "févr.", "mars", "avr.", "mai", "juin",
        "juil.", "août", "sept.", "oct.", "nov.", "déc."
    ]
    return f"{dt.day:02d} {months[dt.month]} {dt.year} {dt.hour:02d}:{dt.minute:02d}"


def get_file_type(ext):
    """Return the preview type for a given extension."""
    ext = ext.lower()
    if ext in PDF_EXTENSIONS:
        return 'pdf'
    if ext in IMAGE_EXTENSIONS:
        return 'image'
    if ext in DOCX_EXTENSIONS:
        return 'docx'
    if ext in RTF_EXTENSIONS:
        return 'rtf'
    return None


def natural_sort_key(name):
    """Generate a sort key that handles numbers naturally.
    'D1' < 'D2' < 'D10' instead of 'D1' < 'D10' < 'D2'
    """
    parts = re.split(r'(\d+)', name.lower())
    result = []
    for part in parts:
        if part.isdigit():
            result.append((1, int(part)))
        else:
            result.append((0, part))
    return result


def get_type_icon_char(file_type, ext):
    """Return an emoji character for file type indication."""
    if file_type == 'pdf':
        return '📄'
    if file_type == 'image':
        return '🖼️'
    if file_type in ('docx', 'rtf'):
        return '📝'
    # Common non-previewable types
    ext = ext.lower()
    if ext in ('.exe', '.msi'):
        return '⚙️'
    if ext in ('.zip', '.rar', '.7z', '.tar', '.gz'):
        return '📦'
    if ext in ('.mp3', '.wav', '.flac', '.aac', '.ogg', '.wma'):
        return '🎵'
    if ext in ('.mp4', '.avi', '.mkv', '.mov', '.wmv', '.flv'):
        return '🎬'
    if ext in ('.py', '.js', '.html', '.css', '.json', '.xml', '.yml', '.yaml',
               '.c', '.cpp', '.h', '.java', '.cs', '.rb', '.go', '.rs', '.sh', '.bat'):
        return '💻'
    if ext in ('.txt', '.md', '.log', '.csv', '.ini', '.cfg'):
        return '📃'
    if ext in ('.xls', '.xlsx', '.ods'):
        return '📊'
    if ext in ('.ppt', '.pptx', '.odp'):
        return '📽️'
    return '📄'


def extract_docx_html(file_path):
    """Extract content from a DOCX file and return HTML."""
    if not HAS_DOCX:
        return "<p style='color:#c45c3e;'>Module python-docx non installé</p>"
    try:
        doc = DocxDocument(file_path)
        html_parts = []
        for para in doc.paragraphs:
            style = para.style.name if para.style else ""
            text = ""
            for run in para.runs:
                t = run.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                if run.bold and run.italic:
                    t = f"<b><i>{t}</i></b>"
                elif run.bold:
                    t = f"<b>{t}</b>"
                elif run.italic:
                    t = f"<i>{t}</i>"
                if run.underline:
                    t = f"<u>{t}</u>"
                text += t

            if "Heading 1" in style:
                html_parts.append(f"<h1>{text}</h1>")
            elif "Heading 2" in style:
                html_parts.append(f"<h2>{text}</h2>")
            elif "Heading 3" in style:
                html_parts.append(f"<h3>{text}</h3>")
            elif "List" in style:
                html_parts.append(f"<p style='margin-left:20px;'>• {text}</p>")
            else:
                html_parts.append(f"<p>{text}</p>" if text.strip() else "<br>")

        # Tables
        for table in doc.tables:
            html_parts.append("<table border='1' cellpadding='6' cellspacing='0' "
                              "style='border-collapse:collapse; border-color:#3a3a42; "
                              "margin:10px 0; width:100%;'>")
            for i, row in enumerate(table.rows):
                html_parts.append("<tr>")
                tag = "th" if i == 0 else "td"
                for cell in row.cells:
                    cell_text = cell.text.replace("&", "&amp;").replace("<", "&lt;")
                    bg = "background-color:#2a2a30;" if i == 0 else ""
                    html_parts.append(
                        f"<{tag} style='{bg} color:#c0bcb4; border:1px solid #3a3a42; "
                        f"padding:6px;'>{cell_text}</{tag}>"
                    )
                html_parts.append("</tr>")
            html_parts.append("</table>")

        return "\n".join(html_parts)
    except Exception as e:
        return f"<p style='color:#c45c3e;'>Erreur de lecture DOCX: {e}</p>"


def extract_rtf_html(file_path):
    """Extract text from RTF file and return HTML."""
    if not HAS_RTF:
        # Fallback: try raw read
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                raw = f.read()
            # Basic RTF stripping if striprtf not available
            text = raw.replace("\\par", "\n").replace("\\tab", "\t")
            paragraphs = text.split("\n")
            html = "\n".join(f"<p>{p}</p>" for p in paragraphs if p.strip())
            return f"<p style='color:#8a8890; font-style:italic;'>Module striprtf non installé — affichage brut</p>{html}"
        except Exception as e:
            return f"<p style='color:#c45c3e;'>Erreur: {e}</p>"

    try:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            rtf_content = f.read()
        text = rtf_to_text(rtf_content)
        paragraphs = text.split("\n")
        html = "\n".join(f"<p>{p.replace('&','&amp;').replace('<','&lt;')}</p>"
                         for p in paragraphs)
        return html
    except Exception as e:
        return f"<p style='color:#c45c3e;'>Erreur de lecture RTF: {e}</p>"


# ─── Stylesheet ───

DARK_STYLE = """
QMainWindow { background-color: #1a1a1e; }

QWidget#toolbar {
    background: qlineargradient(y1:0, y2:1, stop:0 #2a2a30, stop:1 #222228);
    border-bottom: 1px solid #383840;
}
QWidget#toolbar QLabel { color: #8a8890; font-size: 12px; }

QPushButton#openBtn, QPushButton#actionBtn {
    background-color: #c45c3e; color: white; border: none;
    border-radius: 5px; padding: 5px 14px; font-weight: bold; font-size: 12px;
}
QPushButton#openBtn:hover, QPushButton#actionBtn:hover { background-color: #d4714f; }
QPushButton#openBtn:pressed, QPushButton#actionBtn:pressed { background-color: #b04e32; }

QPushButton#zoomBtn {
    background-color: transparent; border: 1px solid #3a3a42; border-radius: 4px;
    color: #b0acaa; padding: 4px 8px; font-size: 13px;
    min-width: 28px; max-width: 28px; min-height: 24px;
}
QPushButton#zoomBtn:hover { background-color: #3a3a42; }

QPushButton#fitBtn {
    background-color: transparent; border: 1px solid #3a3a42; border-radius: 4px;
    color: #b0acaa; padding: 4px 10px; font-size: 11px;
}
QPushButton#fitBtn:hover { background-color: #3a3a42; }
QPushButton#fitBtn[active="true"] { background-color: #3a3a42; color: #e0ddd5; }

QPushButton#navBtn {
    background-color: transparent; border: 1px solid #3a3a42; border-radius: 4px;
    color: #b0acaa; padding: 4px 10px; font-size: 11px; min-width: 32px;
}
QPushButton#navBtn:hover { background-color: #3a3a42; }
QPushButton#navBtn:disabled { color: #3a3a42; border-color: #2a2a30; }

QPushButton#pathBtn {
    background-color: transparent; border: 1px solid #2e2e35; border-radius: 4px;
    color: #a0a0a8; padding: 4px 10px; font-size: 11px; text-align: left;
}
QPushButton#pathBtn:hover { background-color: #2a2a30; border-color: #3a3a42; }

QLineEdit#pathEdit {
    background-color: #22222a; border: 1px solid #c45c3e; border-radius: 4px;
    color: #e0ddd5; padding: 4px 8px; font-size: 12px;
    selection-background-color: #c45c3e;
}

QSplitter { background-color: #1a1a1e; }
QSplitter::handle { background-color: #2e2e35; width: 3px; }
QSplitter::handle:hover { background-color: rgba(196, 92, 62, 0.4); }

QTreeView, QTreeWidget {
    background-color: #1e1e22; border: none; outline: none;
    color: #c0bcb4; font-size: 12px; selection-background-color: transparent;
}
QTreeView::item, QTreeWidget::item {
    padding: 4px 4px; border-left: 3px solid transparent;
}
QTreeView::item:selected, QTreeWidget::item:selected {
    background: qlineargradient(x1:0, x2:1,
        stop:0 rgba(196, 92, 62, 0.15), stop:1 rgba(196, 92, 62, 0.07));
    border-left: 3px solid #c45c3e; color: #f0ece4;
}
QTreeView::item:hover:!selected, QTreeWidget::item:hover:!selected {
    background-color: rgba(255, 255, 255, 0.02);
}

QHeaderView::section {
    background-color: #1c1c20; color: #6a6870; border: none;
    border-bottom: 1px solid #2a2a30; padding: 6px 8px;
    font-size: 10px; font-weight: bold;
}

QStatusBar { background-color: #1c1c20; border-top: 1px solid #2a2a30; color: #5a5860; font-size: 11px; }
QStatusBar QLabel { color: #5a5860; font-size: 11px; padding: 0 4px; }

QMenu {
    background-color: #2a2a30; border: 1px solid #3a3a42; border-radius: 6px;
    padding: 4px 0; color: #c0bcb4; font-size: 12px;
}
QMenu::item { padding: 8px 16px; }
QMenu::item:selected { background-color: rgba(196, 92, 62, 0.2); color: #f0ece4; }
QMenu::separator { height: 1px; background-color: #3a3a42; margin: 4px 8px; }

QPdfView { background-color: #2a2a2e; border: none; }
QLabel#emptyLabel { color: #4a4850; font-size: 14px; background-color: #141416; }
QFrame#previewContainer { background-color: #141416; border: none; }

QTextBrowser {
    background-color: #1e1e22; color: #c0bcb4; border: none;
    font-size: 13px; padding: 16px; selection-background-color: #c45c3e;
}

QScrollArea { background-color: #141416; border: none; }

QScrollBar:vertical { background: transparent; width: 8px; }
QScrollBar::handle:vertical { background: #3a3a42; border-radius: 4px; min-height: 20px; }
QScrollBar::handle:vertical:hover { background: #4a4a52; }
QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical { height: 0; }
QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical { background: transparent; }
QScrollBar:horizontal { background: transparent; height: 8px; }
QScrollBar::handle:horizontal { background: #3a3a42; border-radius: 4px; min-width: 20px; }
QScrollBar::handle:horizontal:hover { background: #4a4a52; }
QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal { width: 0; }
QScrollBar::add-page:horizontal, QScrollBar::sub-page:horizontal { background: transparent; }
"""


# ─── Viewer Window ───

class ViewerWindow(QMainWindow):
    """Fenêtre séparée dédiée à la prévisualisation multi-format."""

    def __init__(self):
        super().__init__()
        self.setWindowTitle("PDF Quick Browser — Preview")
        self.setMinimumSize(700, 500)
        self.resize(900, 900)

        self.doc_cache = {}
        self.cache_size = 10
        self.active_doc = None
        self.active_path = None
        self.active_type = None
        self.explorer = None  # set by ExplorerWindow after init

        self._setup_ui()
        self._setup_shortcuts()
        self.setStyleSheet(DARK_STYLE)

    def _setup_ui(self):
        central = QWidget()
        self.setCentralWidget(central)
        layout = QVBoxLayout(central)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        # Toolbar
        tb = QWidget()
        tb.setObjectName("toolbar")
        tb.setFixedHeight(44)
        tl = QHBoxLayout(tb)
        tl.setContentsMargins(12, 0, 12, 0)
        tl.setSpacing(8)

        self.file_label = QLabel("  Aucun fichier sélectionné")
        self.file_label.setStyleSheet("color: #b0acaa; font-size: 12px;")
        tl.addWidget(self.file_label)
        tl.addStretch()

        # Page nav (PDF only)
        self.btn_page_up = QPushButton("▲")
        self.btn_page_up.setObjectName("navBtn")
        self.btn_page_up.setCursor(Qt.CursorShape.PointingHandCursor)
        self.btn_page_up.clicked.connect(self._page_up)
        tl.addWidget(self.btn_page_up)

        self.page_label = QLabel("")
        self.page_label.setFixedWidth(70)
        self.page_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.page_label.setStyleSheet("color: #8a8890; font-size: 11px;")
        tl.addWidget(self.page_label)

        self.btn_page_down = QPushButton("▼")
        self.btn_page_down.setObjectName("navBtn")
        self.btn_page_down.setCursor(Qt.CursorShape.PointingHandCursor)
        self.btn_page_down.clicked.connect(self._page_down)
        tl.addWidget(self.btn_page_down)

        self.sep_nav = QLabel("  │  ")
        self.sep_nav.setStyleSheet("color: #3a3a42;")
        tl.addWidget(self.sep_nav)

        # Zoom
        self.btn_fit_width = QPushButton("Largeur")
        self.btn_fit_width.setObjectName("fitBtn")
        self.btn_fit_width.setProperty("active", True)
        self.btn_fit_width.setCursor(Qt.CursorShape.PointingHandCursor)
        self.btn_fit_width.clicked.connect(self._zoom_fit_width)
        tl.addWidget(self.btn_fit_width)

        self.btn_fit_page = QPushButton("Page")
        self.btn_fit_page.setObjectName("fitBtn")
        self.btn_fit_page.setCursor(Qt.CursorShape.PointingHandCursor)
        self.btn_fit_page.clicked.connect(self._zoom_fit_page)
        tl.addWidget(self.btn_fit_page)

        self.btn_zoom_out = QPushButton("−")
        self.btn_zoom_out.setObjectName("zoomBtn")
        self.btn_zoom_out.setCursor(Qt.CursorShape.PointingHandCursor)
        self.btn_zoom_out.clicked.connect(lambda: self._adjust_zoom(-0.25))
        tl.addWidget(self.btn_zoom_out)

        self.zoom_label = QLabel("Auto")
        self.zoom_label.setFixedWidth(50)
        self.zoom_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.zoom_label.setStyleSheet("color: #8a8890; font-size: 11px;")
        tl.addWidget(self.zoom_label)

        self.btn_zoom_in = QPushButton("+")
        self.btn_zoom_in.setObjectName("zoomBtn")
        self.btn_zoom_in.setCursor(Qt.CursorShape.PointingHandCursor)
        self.btn_zoom_in.clicked.connect(lambda: self._adjust_zoom(0.25))
        tl.addWidget(self.btn_zoom_in)

        layout.addWidget(tb)

        # ── Stacked preview widgets ──
        container = QFrame()
        container.setObjectName("previewContainer")
        cl = QVBoxLayout(container)
        cl.setContentsMargins(0, 0, 0, 0)
        cl.setSpacing(0)

        self.stack = QStackedWidget()

        # Page 0: Empty state
        self.empty_label = QLabel(
            "📄\n\n"
            "Sélectionnez un fichier dans l'explorateur\n"
            "pour le prévisualiser ici\n\n"
            "Formats supportés : PDF, Images, DOCX, RTF"
        )
        self.empty_label.setObjectName("emptyLabel")
        self.empty_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.stack.addWidget(self.empty_label)  # index 0

        # Page 1: PDF viewer
        self.pdf_view = QPdfView()
        self.pdf_view.setPageMode(QPdfView.PageMode.MultiPage)
        self.pdf_view.setZoomMode(QPdfView.ZoomMode.FitToWidth)
        self.stack.addWidget(self.pdf_view)  # index 1

        # Page 2: Image viewer
        self.image_scroll = QScrollArea()
        self.image_scroll.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.image_scroll.setWidgetResizable(True)
        self.image_label = QLabel()
        self.image_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.image_label.setStyleSheet("background-color: #141416;")
        self.image_scroll.setWidget(self.image_label)
        self.stack.addWidget(self.image_scroll)  # index 2

        # Page 3: Document viewer (DOCX/RTF)
        self.doc_browser = QTextBrowser()
        self.doc_browser.setOpenExternalLinks(False)
        self.doc_browser.setStyleSheet("""
            QTextBrowser {
                background-color: #1e1e22; color: #c0bcb4; border: none;
                font-family: 'Segoe UI', sans-serif; font-size: 13px;
                padding: 24px; selection-background-color: #c45c3e;
            }
        """)
        self.stack.addWidget(self.doc_browser)  # index 3

        cl.addWidget(self.stack)
        layout.addWidget(container)

        # Status
        self.statusBar().show()
        self.status_info = QLabel("")
        self.statusBar().addWidget(self.status_info, 1)

        # Start with empty state
        self.stack.setCurrentIndex(0)
        self._set_pdf_controls_visible(False)

    def _setup_shortcuts(self):
        QShortcut(QKeySequence("+"), self, lambda: self._adjust_zoom(0.25))
        QShortcut(QKeySequence("="), self, lambda: self._adjust_zoom(0.25))
        QShortcut(QKeySequence("-"), self, lambda: self._adjust_zoom(-0.25))
        QShortcut(QKeySequence("Ctrl+0"), self, self._zoom_fit_width)

    def _set_pdf_controls_visible(self, visible):
        """Show/hide PDF-specific controls."""
        for w in [self.btn_page_up, self.btn_page_down, self.page_label,
                  self.sep_nav, self.btn_fit_width, self.btn_fit_page]:
            w.setVisible(visible)

    def _set_zoom_visible(self, visible):
        """Show/hide zoom controls."""
        for w in [self.btn_zoom_in, self.btn_zoom_out, self.zoom_label]:
            w.setVisible(visible)

    # ── Public API ──

    def show_file(self, file_path):
        """Show any supported file."""
        if file_path == self.active_path:
            return

        ext = os.path.splitext(file_path)[1].lower()
        file_type = get_file_type(ext)

        if file_type == 'pdf':
            self._show_pdf(file_path)
        elif file_type == 'image':
            self._show_image(file_path)
        elif file_type == 'docx':
            self._show_docx(file_path)
        elif file_type == 'rtf':
            self._show_rtf(file_path)
        else:
            return

        self.active_path = file_path
        self.active_type = file_type

    def preload_pdf(self, file_path):
        """Preload a PDF into cache."""
        if file_path in self.doc_cache or len(self.doc_cache) >= self.cache_size:
            return
        doc = QPdfDocument(self)
        if doc.load(file_path) == QPdfDocument.Error.None_:
            self.doc_cache[file_path] = doc

    def clear_cache(self):
        for doc in self.doc_cache.values():
            doc.close()
        self.doc_cache.clear()
        self.active_doc = None
        self.active_path = None

    # ── PDF ──

    def _show_pdf(self, file_path):
        self._set_pdf_controls_visible(True)
        self._set_zoom_visible(True)

        if file_path in self.doc_cache:
            doc = self.doc_cache[file_path]
        else:
            doc = QPdfDocument(self)
            if doc.load(file_path) != QPdfDocument.Error.None_:
                self.status_info.setText(f"  ✕ Erreur PDF: {file_path}")
                return
            if len(self.doc_cache) >= self.cache_size:
                oldest = next(iter(self.doc_cache))
                if oldest != self.active_path:
                    self.doc_cache.pop(oldest).close()
            self.doc_cache[file_path] = doc

        self.active_doc = doc
        self.pdf_view.setDocument(doc)
        self.stack.setCurrentIndex(1)

        # Reset to first page
        nav = self.pdf_view.pageNavigator()
        nav.jump(0, QPointF())

        name = os.path.basename(file_path)
        pc = doc.pageCount()
        self.file_label.setText(f"  📄 {name}")
        self.setWindowTitle(f"Preview — {name}")
        self.status_info.setText(
            f"  {name}  —  {pc} page{'s' if pc != 1 else ''}  —  "
            f"{format_size(os.path.getsize(file_path))}"
        )
        self._update_page_info()

    # ── Image ──

    def _show_image(self, file_path):
        self._set_pdf_controls_visible(False)
        self._set_zoom_visible(True)
        self.active_doc = None

        pixmap = QPixmap(file_path)
        if pixmap.isNull():
            self.status_info.setText(f"  ✕ Erreur image: {file_path}")
            return

        # Scale to fit scroll area while keeping aspect ratio
        self._current_pixmap = pixmap
        self._update_image_display()
        self.stack.setCurrentIndex(2)

        name = os.path.basename(file_path)
        w, h = pixmap.width(), pixmap.height()
        self.file_label.setText(f"  🖼️ {name}")
        self.setWindowTitle(f"Preview — {name}")
        self.status_info.setText(
            f"  {name}  —  {w}×{h} px  —  "
            f"{format_size(os.path.getsize(file_path))}"
        )

    def _update_image_display(self):
        """Update image display with current zoom."""
        if not hasattr(self, '_current_pixmap'):
            return
        pixmap = self._current_pixmap
        available = self.image_scroll.size()
        if self.pdf_view.zoomMode() != QPdfView.ZoomMode.Custom:
            # Fit to window
            scaled = pixmap.scaled(
                available.width() - 20, available.height() - 20,
                Qt.AspectRatioMode.KeepAspectRatio,
                Qt.TransformationMode.SmoothTransformation
            )
        else:
            factor = self.pdf_view.zoomFactor()
            scaled = pixmap.scaled(
                int(pixmap.width() * factor), int(pixmap.height() * factor),
                Qt.AspectRatioMode.KeepAspectRatio,
                Qt.TransformationMode.SmoothTransformation
            )
        self.image_label.setPixmap(scaled)

    # ── DOCX ──

    def _show_docx(self, file_path):
        self._set_pdf_controls_visible(False)
        self._set_zoom_visible(False)
        self.active_doc = None

        html = extract_docx_html(file_path)
        self.doc_browser.setHtml(f"""
            <html><body style="background-color:#1e1e22; color:#c0bcb4;
                font-family:'Segoe UI',sans-serif; font-size:13px; padding:8px;">
            <h2 style="color:#f0ece4; border-bottom:1px solid #3a3a42;
                padding-bottom:8px;">{os.path.basename(file_path)}</h2>
            {html}
            </body></html>
        """)
        self.stack.setCurrentIndex(3)

        name = os.path.basename(file_path)
        self.file_label.setText(f"  📝 {name}")
        self.setWindowTitle(f"Preview — {name}")
        self.status_info.setText(
            f"  {name}  —  DOCX  —  {format_size(os.path.getsize(file_path))}"
        )

    # ── RTF ──

    def _show_rtf(self, file_path):
        self._set_pdf_controls_visible(False)
        self._set_zoom_visible(False)
        self.active_doc = None

        html = extract_rtf_html(file_path)
        self.doc_browser.setHtml(f"""
            <html><body style="background-color:#1e1e22; color:#c0bcb4;
                font-family:'Segoe UI',sans-serif; font-size:13px; padding:8px;">
            <h2 style="color:#f0ece4; border-bottom:1px solid #3a3a42;
                padding-bottom:8px;">{os.path.basename(file_path)}</h2>
            {html}
            </body></html>
        """)
        self.stack.setCurrentIndex(3)

        name = os.path.basename(file_path)
        self.file_label.setText(f"  📝 {name}")
        self.setWindowTitle(f"Preview — {name}")
        self.status_info.setText(
            f"  {name}  —  RTF  —  {format_size(os.path.getsize(file_path))}"
        )

    # ── PDF page nav ──

    def _page_up(self):
        nav = self.pdf_view.pageNavigator()
        if nav.currentPage() > 0:
            nav.jump(nav.currentPage() - 1, nav.currentLocation())
            self._update_page_info()

    def _page_down(self):
        nav = self.pdf_view.pageNavigator()
        if self.active_doc and nav.currentPage() < self.active_doc.pageCount() - 1:
            nav.jump(nav.currentPage() + 1, nav.currentLocation())
            self._update_page_info()

    def _update_page_info(self):
        if self.active_doc and self.active_doc.pageCount() > 0:
            nav = self.pdf_view.pageNavigator()
            cp = nav.currentPage() + 1
            total = self.active_doc.pageCount()
            self.page_label.setText(f"p. {cp}/{total}")
            self.btn_page_up.setEnabled(cp > 1)
            self.btn_page_down.setEnabled(cp < total)
        else:
            self.page_label.setText("")

    # ── Zoom ──

    def _adjust_zoom(self, delta):
        if self.active_type == 'image':
            current = self.pdf_view.zoomFactor() if self.pdf_view.zoomMode() == QPdfView.ZoomMode.Custom else 1.0
            new_zoom = max(0.1, min(5.0, current + delta))
            self.pdf_view.setZoomMode(QPdfView.ZoomMode.Custom)
            self.pdf_view.setZoomFactor(new_zoom)
            self.zoom_label.setText(f"{int(new_zoom * 100)}%")
            self._update_fit_buttons(None)
            self._update_image_display()
        elif self.active_type == 'pdf':
            new_zoom = max(0.25, min(5.0, self.pdf_view.zoomFactor() + delta))
            self.pdf_view.setZoomMode(QPdfView.ZoomMode.Custom)
            self.pdf_view.setZoomFactor(new_zoom)
            self.zoom_label.setText(f"{int(new_zoom * 100)}%")
            self._update_fit_buttons(None)

    def _zoom_fit_width(self):
        self.pdf_view.setZoomMode(QPdfView.ZoomMode.FitToWidth)
        self.zoom_label.setText("Auto")
        self._update_fit_buttons("width")
        if self.active_type == 'image':
            self._update_image_display()

    def _zoom_fit_page(self):
        self.pdf_view.setZoomMode(QPdfView.ZoomMode.FitInView)
        self.zoom_label.setText("Auto")
        self._update_fit_buttons("page")
        if self.active_type == 'image':
            self._update_image_display()

    def _update_fit_buttons(self, active):
        self.btn_fit_width.setProperty("active", active == "width")
        self.btn_fit_page.setProperty("active", active == "page")
        for btn in [self.btn_fit_width, self.btn_fit_page]:
            btn.style().unpolish(btn)
            btn.style().polish(btn)

    def resizeEvent(self, event):
        super().resizeEvent(event)
        if self.active_type == 'image' and self.pdf_view.zoomMode() != QPdfView.ZoomMode.Custom:
            QTimer.singleShot(10, self._update_image_display)

    def keyPressEvent(self, event):
        """Forward ↑↓ navigation to explorer file list."""
        key = event.key()
        if key in (Qt.Key.Key_Up, Qt.Key.Key_Down) and self.explorer:
            fl = self.explorer.file_list
            current = fl.currentItem()
            if current is None:
                return
            idx = fl.indexOfTopLevelItem(current)
            if key == Qt.Key.Key_Down:
                new_idx = min(idx + 1, fl.topLevelItemCount() - 1)
            else:
                new_idx = max(idx - 1, 0)
            if new_idx != idx:
                fl.setCurrentItem(fl.topLevelItem(new_idx))
            return
        if key == Qt.Key.Key_PageDown and self.explorer:
            fl = self.explorer.file_list
            current = fl.currentItem()
            if current:
                idx = fl.indexOfTopLevelItem(current)
                new_idx = min(idx + 10, fl.topLevelItemCount() - 1)
                fl.setCurrentItem(fl.topLevelItem(new_idx))
            return
        if key == Qt.Key.Key_PageUp and self.explorer:
            fl = self.explorer.file_list
            current = fl.currentItem()
            if current:
                idx = fl.indexOfTopLevelItem(current)
                new_idx = max(idx - 10, 0)
                fl.setCurrentItem(fl.topLevelItem(new_idx))
            return
        if key == Qt.Key.Key_Home and self.explorer:
            fl = self.explorer.file_list
            if fl.topLevelItemCount() > 0:
                fl.setCurrentItem(fl.topLevelItem(0))
            return
        if key == Qt.Key.Key_End and self.explorer:
            fl = self.explorer.file_list
            if fl.topLevelItemCount() > 0:
                fl.setCurrentItem(fl.topLevelItem(fl.topLevelItemCount() - 1))
            return
        super().keyPressEvent(event)

    def closeEvent(self, event):
        self.hide()
        event.ignore()


# ─── Explorer Window ───

class ExplorerWindow(QMainWindow):
    """Fenêtre principale avec explorateur de fichiers complet."""

    def __init__(self, viewer: ViewerWindow):
        super().__init__()
        self.viewer = viewer
        self.viewer.explorer = self  # allow viewer to navigate files
        self.setWindowTitle("PDF Quick Browser — Explorateur")
        self.setMinimumSize(500, 400)
        self.resize(750, 900)

        self.current_path = ""
        self.history = []
        self.history_pos = -1

        self._setup_ui()
        self._setup_shortcuts()
        self.setStyleSheet(DARK_STYLE)

        home = os.path.expanduser("~")
        self._navigate_to(home)

    def _setup_ui(self):
        central = QWidget()
        self.setCentralWidget(central)
        main_layout = QVBoxLayout(central)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Navigation toolbar
        nav_bar = QWidget()
        nav_bar.setObjectName("toolbar")
        nav_bar.setFixedHeight(44)
        nl = QHBoxLayout(nav_bar)
        nl.setContentsMargins(8, 0, 8, 0)
        nl.setSpacing(4)

        self.btn_back = QPushButton("◀")
        self.btn_back.setObjectName("navBtn")
        self.btn_back.setToolTip("Précédent (Alt+←)")
        self.btn_back.setCursor(Qt.CursorShape.PointingHandCursor)
        self.btn_back.clicked.connect(self._go_back)
        self.btn_back.setEnabled(False)
        nl.addWidget(self.btn_back)

        self.btn_forward = QPushButton("▶")
        self.btn_forward.setObjectName("navBtn")
        self.btn_forward.setToolTip("Suivant (Alt+→)")
        self.btn_forward.setCursor(Qt.CursorShape.PointingHandCursor)
        self.btn_forward.clicked.connect(self._go_forward)
        self.btn_forward.setEnabled(False)
        nl.addWidget(self.btn_forward)

        self.btn_up = QPushButton("▲")
        self.btn_up.setObjectName("navBtn")
        self.btn_up.setToolTip("Dossier parent (Alt+↑)")
        self.btn_up.setCursor(Qt.CursorShape.PointingHandCursor)
        self.btn_up.clicked.connect(self._go_up)
        nl.addWidget(self.btn_up)

        self.btn_home = QPushButton("🏠")
        self.btn_home.setObjectName("navBtn")
        self.btn_home.setToolTip("Dossier personnel")
        self.btn_home.setCursor(Qt.CursorShape.PointingHandCursor)
        self.btn_home.clicked.connect(lambda: self._navigate_to(os.path.expanduser("~")))
        nl.addWidget(self.btn_home)

        nl.addSpacing(4)

        self.path_button = QPushButton("")
        self.path_button.setObjectName("pathBtn")
        self.path_button.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Fixed)
        self.path_button.setCursor(Qt.CursorShape.IBeamCursor)
        self.path_button.clicked.connect(self._start_path_edit)
        nl.addWidget(self.path_button)

        self.path_edit = QLineEdit()
        self.path_edit.setObjectName("pathEdit")
        self.path_edit.setVisible(False)
        self.path_edit.returnPressed.connect(self._finish_path_edit)
        nl.addWidget(self.path_edit)

        nl.addSpacing(4)

        self.btn_show_viewer = QPushButton("📄 Viewer")
        self.btn_show_viewer.setObjectName("actionBtn")
        self.btn_show_viewer.setCursor(Qt.CursorShape.PointingHandCursor)
        self.btn_show_viewer.setToolTip("Afficher la fenêtre de preview")
        self.btn_show_viewer.clicked.connect(self._show_viewer)
        nl.addWidget(self.btn_show_viewer)

        main_layout.addWidget(nav_bar)

        # Splitter: tree + file list
        self.splitter = QSplitter(Qt.Orientation.Horizontal)

        self.fs_model = QFileSystemModel()
        self.fs_model.setRootPath("")
        self.fs_model.setFilter(QDir.Filter.AllDirs | QDir.Filter.NoDotAndDotDot | QDir.Filter.Drives)

        self.folder_tree = QTreeView()
        self.folder_tree.setModel(self.fs_model)
        self.folder_tree.setRootIndex(self.fs_model.index(""))
        self.folder_tree.setHeaderHidden(True)
        for i in range(1, self.fs_model.columnCount()):
            self.folder_tree.hideColumn(i)
        self.folder_tree.setMinimumWidth(180)
        self.folder_tree.setAnimated(True)
        self.folder_tree.clicked.connect(self._on_folder_clicked)

        self.file_list = QTreeWidget()
        self.file_list.setHeaderLabels(["Nom", "Taille", "Date", "Type"])
        self.file_list.setRootIsDecorated(False)
        self.file_list.setSelectionMode(QAbstractItemView.SelectionMode.SingleSelection)
        self.file_list.setFocusPolicy(Qt.FocusPolicy.StrongFocus)
        self.file_list.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.file_list.customContextMenuRequested.connect(self._show_context_menu)
        self.file_list.currentItemChanged.connect(self._on_file_selection_changed)
        self.file_list.itemDoubleClicked.connect(self._on_file_double_clicked)
        self.file_list.setSortingEnabled(False)

        header = self.file_list.header()
        header.setStretchLastSection(False)
        header.setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
        header.setSectionResizeMode(1, QHeaderView.ResizeMode.Fixed)
        header.setSectionResizeMode(2, QHeaderView.ResizeMode.Fixed)
        header.setSectionResizeMode(3, QHeaderView.ResizeMode.Fixed)
        header.resizeSection(1, 80)
        header.resizeSection(2, 140)
        header.resizeSection(3, 60)

        self.splitter.addWidget(self.folder_tree)
        self.splitter.addWidget(self.file_list)
        self.splitter.setSizes([220, 530])
        self.splitter.setCollapsible(0, False)
        self.splitter.setCollapsible(1, False)

        main_layout.addWidget(self.splitter)

        self.statusBar().show()
        self.status_path = QLabel("")
        self.status_count = QLabel("")
        self.statusBar().addWidget(self.status_path, 1)
        self.statusBar().addPermanentWidget(self.status_count)

    def _setup_shortcuts(self):
        QShortcut(QKeySequence("Alt+Left"), self, self._go_back)
        QShortcut(QKeySequence("Alt+Right"), self, self._go_forward)
        QShortcut(QKeySequence("Alt+Up"), self, self._go_up)
        QShortcut(QKeySequence("Ctrl+L"), self, self._start_path_edit)
        QShortcut(QKeySequence("Ctrl+E"), self, self._open_in_explorer)
        QShortcut(QKeySequence("F5"), self, self._refresh)
        QShortcut(QKeySequence("Backspace"), self, self._go_up)
        QShortcut(QKeySequence("Return"), self, self._enter_selected)
        QShortcut(QKeySequence("Escape"), self, self._cancel_path_edit)

    # Navigation

    def _navigate_to(self, path, add_history=True):
        path = os.path.normpath(path)
        if not os.path.isdir(path):
            return

        self.current_path = path

        if add_history:
            if self.history_pos < len(self.history) - 1:
                self.history = self.history[:self.history_pos + 1]
            self.history.append(path)
            self.history_pos = len(self.history) - 1

        self._update_nav_buttons()
        self._update_path_bar()
        self._populate_file_list()

        idx = self.fs_model.index(path)
        if idx.isValid():
            self.folder_tree.setCurrentIndex(idx)
            self.folder_tree.scrollTo(idx)
            parent = idx.parent()
            while parent.isValid():
                self.folder_tree.expand(parent)
                parent = parent.parent()

    def _go_back(self):
        if self.history_pos > 0:
            self.history_pos -= 1
            self._navigate_to(self.history[self.history_pos], add_history=False)

    def _go_forward(self):
        if self.history_pos < len(self.history) - 1:
            self.history_pos += 1
            self._navigate_to(self.history[self.history_pos], add_history=False)

    def _go_up(self):
        parent = os.path.dirname(self.current_path)
        if parent and parent != self.current_path:
            self._navigate_to(parent)

    def _update_nav_buttons(self):
        self.btn_back.setEnabled(self.history_pos > 0)
        self.btn_forward.setEnabled(self.history_pos < len(self.history) - 1)
        parent = os.path.dirname(self.current_path)
        self.btn_up.setEnabled(bool(parent) and parent != self.current_path)

    def _update_path_bar(self):
        self.path_button.setText(f"  {self.current_path}")
        self.status_path.setText(f"  {self.current_path}")

    def _start_path_edit(self):
        self.path_button.setVisible(False)
        self.path_edit.setVisible(True)
        self.path_edit.setText(self.current_path)
        self.path_edit.setFocus()
        self.path_edit.selectAll()

    def _finish_path_edit(self):
        new_path = self.path_edit.text().strip()
        self.path_edit.setVisible(False)
        self.path_button.setVisible(True)
        if os.path.isdir(new_path):
            self._navigate_to(new_path)

    def _cancel_path_edit(self):
        if self.path_edit.isVisible():
            self.path_edit.setVisible(False)
            self.path_button.setVisible(True)

    # File list

    def _populate_file_list(self):
        self.file_list.setSortingEnabled(False)
        self.file_list.clear()
        entries_dirs = []
        entries_files = []

        try:
            for entry in os.scandir(self.current_path):
                try:
                    stat = entry.stat()
                    is_dir = entry.is_dir(follow_symlinks=False)
                    if is_dir:
                        entries_dirs.append((entry.name, entry.path, stat))
                    else:
                        entries_files.append((entry.name, entry.path, stat))
                except (PermissionError, OSError):
                    continue
        except (PermissionError, OSError):
            pass

        # Natural sort
        entries_dirs.sort(key=lambda e: natural_sort_key(e[0]))
        entries_files.sort(key=lambda e: natural_sort_key(e[0]))

        # Add folders first
        for name, path, stat in entries_dirs:
            item = QTreeWidgetItem([
                f"📁  {name}", "", format_date(stat.st_mtime), "Dossier"
            ])
            item.setData(0, Qt.ItemDataRole.UserRole, {
                'path': path, 'is_dir': True, 'name': name,
            })
            self.file_list.addTopLevelItem(item)

        # Then files
        preview_count = 0
        for name, path, stat in entries_files:
            ext = os.path.splitext(name)[1].lower()
            file_type = get_file_type(ext)
            icon_char = get_type_icon_char(file_type, ext)

            if file_type is not None:
                preview_count += 1

            item = QTreeWidgetItem([
                f"{icon_char}  {name}", format_size(stat.st_size),
                format_date(stat.st_mtime),
                ext[1:].upper() if ext else ""
            ])
            item.setData(0, Qt.ItemDataRole.UserRole, {
                'path': path, 'is_dir': False, 'name': name,
                'previewable': file_type is not None,
                'file_type': file_type,
            })
            self.file_list.addTopLevelItem(item)

        self.status_count.setText(
            f"{len(entries_dirs)} dossier{'s' if len(entries_dirs) != 1 else ''}  |  "
            f"{len(entries_files)} fichier{'s' if len(entries_files) != 1 else ''}  "
            f"({preview_count} prévisualisable{'s' if preview_count != 1 else ''})  "
        )

    def _on_folder_clicked(self, index: QModelIndex):
        path = self.fs_model.filePath(index)
        if path and os.path.isdir(path):
            self._navigate_to(path)

    def _on_file_selection_changed(self, current, previous):
        if current is None:
            return
        data = current.data(0, Qt.ItemDataRole.UserRole)
        if not data:
            return

        if not data.get('is_dir') and data.get('previewable'):
            self.viewer.show_file(data['path'])
            QTimer.singleShot(10, lambda: self._preload_adjacent(current))

    def _preload_adjacent(self, current_item):
        idx = self.file_list.indexOfTopLevelItem(current_item)
        for offset in [-1, 1, -2, 2]:
            i = idx + offset
            if 0 <= i < self.file_list.topLevelItemCount():
                item = self.file_list.topLevelItem(i)
                data = item.data(0, Qt.ItemDataRole.UserRole)
                if data and data.get('file_type') == 'pdf':
                    self.viewer.preload_pdf(data['path'])

    def _on_file_double_clicked(self, item, column):
        data = item.data(0, Qt.ItemDataRole.UserRole)
        if data and data.get('is_dir'):
            self._navigate_to(data['path'])

    def _enter_selected(self):
        item = self.file_list.currentItem()
        if item:
            data = item.data(0, Qt.ItemDataRole.UserRole)
            if data and data.get('is_dir'):
                self._navigate_to(data['path'])

    def _refresh(self):
        self._populate_file_list()

    def _show_context_menu(self, position):
        item = self.file_list.itemAt(position)
        if not item:
            return
        data = item.data(0, Qt.ItemDataRole.UserRole)
        if not data:
            return

        menu = QMenu(self)
        if data.get('is_dir'):
            menu.addAction("📂  Ouvrir", lambda: self._navigate_to(data['path']))
            menu.addSeparator()
        menu.addAction("📂  Ouvrir dans l'Explorateur",
                       lambda: os.system(f'explorer /select,"{data["path"]}"'))
        menu.addAction("📋  Copier le chemin",
                       lambda: QApplication.clipboard().setText(data['path']))
        menu.addAction("📝  Copier le nom",
                       lambda: QApplication.clipboard().setText(data['name']))
        menu.exec(self.file_list.viewport().mapToGlobal(position))

    def _open_in_explorer(self):
        item = self.file_list.currentItem()
        if item:
            data = item.data(0, Qt.ItemDataRole.UserRole)
            if data:
                os.system(f'explorer /select,"{data["path"]}"')

    def _show_viewer(self):
        self.viewer.show()
        self.viewer.raise_()
        self.viewer.activateWindow()

    def keyPressEvent(self, event):
        key = event.key()
        no_mod = event.modifiers() == Qt.KeyboardModifier.NoModifier

        if key == Qt.Key.Key_PageDown and no_mod and self.file_list.hasFocus():
            idx = self.file_list.indexOfTopLevelItem(self.file_list.currentItem())
            new_idx = min(idx + 10, self.file_list.topLevelItemCount() - 1)
            if new_idx >= 0:
                self.file_list.setCurrentItem(self.file_list.topLevelItem(new_idx))
            return
        if key == Qt.Key.Key_PageUp and no_mod and self.file_list.hasFocus():
            idx = self.file_list.indexOfTopLevelItem(self.file_list.currentItem())
            new_idx = max(idx - 10, 0)
            self.file_list.setCurrentItem(self.file_list.topLevelItem(new_idx))
            return

        super().keyPressEvent(event)

    def closeEvent(self, event):
        self.viewer.clear_cache()
        self.viewer.close()
        self.viewer.deleteLater()
        event.accept()
        QApplication.quit()


# ─── Main ───

def main():
    app = QApplication(sys.argv)
    app.setStyle("Fusion")

    palette = QPalette()
    palette.setColor(QPalette.ColorRole.Window, QColor("#1a1a1e"))
    palette.setColor(QPalette.ColorRole.WindowText, QColor("#e0ddd5"))
    palette.setColor(QPalette.ColorRole.Base, QColor("#1e1e22"))
    palette.setColor(QPalette.ColorRole.AlternateBase, QColor("#242428"))
    palette.setColor(QPalette.ColorRole.Text, QColor("#e0ddd5"))
    palette.setColor(QPalette.ColorRole.Button, QColor("#2a2a30"))
    palette.setColor(QPalette.ColorRole.ButtonText, QColor("#e0ddd5"))
    palette.setColor(QPalette.ColorRole.Highlight, QColor("#c45c3e"))
    palette.setColor(QPalette.ColorRole.HighlightedText, QColor("#ffffff"))
    app.setPalette(palette)
    app.setFont(QFont("Segoe UI", 10))

    viewer = ViewerWindow()
    explorer = ExplorerWindow(viewer)

    # Set window icon (works both in dev and PyInstaller bundle)
    if getattr(sys, 'frozen', False):
        base_path = sys._MEIPASS
    else:
        base_path = os.path.dirname(os.path.abspath(__file__))
    icon_path = os.path.join(base_path, 'icon.ico')
    if os.path.exists(icon_path):
        app_icon = QIcon(icon_path)
        app.setWindowIcon(app_icon)
        explorer.setWindowIcon(app_icon)
        viewer.setWindowIcon(app_icon)

    screen = app.primaryScreen().availableGeometry()
    explorer_w = int(screen.width() * 0.38)
    viewer_w = screen.width() - explorer_w - 10

    explorer.resize(explorer_w, screen.height() - 40)
    explorer.move(screen.x(), screen.y())

    viewer.resize(viewer_w, screen.height() - 40)
    viewer.move(screen.x() + explorer_w + 5, screen.y())

    explorer.show()
    viewer.show()

    sys.exit(app.exec())


if __name__ == "__main__":
    main()
